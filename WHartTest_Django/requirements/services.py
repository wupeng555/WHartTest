import logging
import json
import re
from string import Template
from typing import List, Dict, Any, Optional
from django.conf import settings
from langchain_openai import ChatOpenAI
from langchain_anthropic import ChatAnthropic
from langchain_core.messages import HumanMessage, SystemMessage
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langgraph_integration.models import LLMConfig
from .models import RequirementDocument, RequirementModule
from prompts.models import UserPrompt

logger = logging.getLogger(__name__)


def create_llm_instance(active_config, temperature=0.1):
    """
    根据配置创建合适的LLM实例，支持多种供应商
    """
    model_identifier = active_config.name or "gpt-3.5-turbo"
    
    # 检测供应商类型
    api_url = active_config.api_url.lower()
    
    if "anthropic.com" in api_url or "claude" in model_identifier.lower():
        # Anthropic/Claude
        llm = ChatAnthropic(
            model=model_identifier,
            api_key=active_config.api_key,
            temperature=temperature
        )
        logger.info(f"Initialized ChatAnthropic with model: {model_identifier}")
    else:
        # OpenAI 兼容 (包括 OpenAI, Ollama, 其他兼容服务)
        llm_kwargs = {
            "model": model_identifier,
            "temperature": temperature,
            "api_key": active_config.api_key,
        }
        
        # 如果不是官方OpenAI，设置base_url
        if "api.openai.com" not in api_url:
            llm_kwargs["base_url"] = active_config.api_url
            # Ollama通常不需要真实的API key
            if "ollama" in api_url or not active_config.api_key:
                llm_kwargs["api_key"] = "ollama"
        
        llm = ChatOpenAI(**llm_kwargs)
        logger.info(f"Initialized ChatOpenAI with model: {model_identifier}")
    
    return llm


class DocumentProcessor:
    """文档处理器 - 负责文档内容提取和预处理"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,
            chunk_overlap=200,
            separators=["\n\n", "\n", "。", ".", " ", ""]
        )
    
    def extract_content(self, document: RequirementDocument) -> str:
        """提取文档内容"""
        try:
            if document.content:
                return document.content
            
            if document.file:
                # TODO: 根据文件类型提取内容
                # 这里先返回模拟内容
                return self._extract_from_file(document.file)
            
            return ""
        except Exception as e:
            logger.error(f"提取文档内容失败: {e}")
            return ""
    
    def _extract_from_file(self, file) -> str:
        """从文件提取内容"""
        try:
            file_extension = file.name.lower().split('.')[-1] if '.' in file.name else ''

            if file_extension == 'txt':
                return self._extract_from_txt(file)
            elif file_extension == 'md':
                return self._extract_from_markdown(file)
            elif file_extension in ['pdf']:
                return self._extract_from_pdf(file)
            elif file_extension in ['docx', 'doc']:
                return self._extract_from_word(file)
            else:
                # 对于不支持的格式，尝试作为文本文件读取
                logger.warning(f"不支持的文件格式: {file_extension}，尝试作为文本文件读取")
                return self._extract_from_txt(file)

        except Exception as e:
            logger.error(f"文件内容提取失败: {e}")
            return ""

    def _extract_from_txt(self, file) -> str:
        """提取TXT文件内容"""
        try:
            # 重置文件指针到开始位置
            file.seek(0)
            content = file.read()

            if isinstance(content, bytes):
                # 尝试不同的编码
                for encoding in ['utf-8', 'gbk', 'gb2312']:
                    try:
                        decoded_content = content.decode(encoding)
                        logger.info(f"成功使用 {encoding} 编码读取文件，内容长度: {len(decoded_content)}")
                        return decoded_content
                    except UnicodeDecodeError:
                        continue
                # 如果都失败了，使用错误处理
                decoded_content = content.decode('utf-8', errors='ignore')
                logger.warning(f"使用UTF-8错误忽略模式读取文件，内容长度: {len(decoded_content)}")
                return decoded_content
            else:
                logger.info(f"直接读取文本内容，长度: {len(content)}")
                return content
        except Exception as e:
            logger.error(f"TXT文件读取失败: {e}")
            return ""

    def _extract_from_markdown(self, file) -> str:
        """提取Markdown文件内容"""
        return self._extract_from_txt(file)  # Markdown本质上是文本文件

    def _extract_from_pdf(self, file) -> str:
        """提取PDF文件内容"""
        try:
            from pypdf import PdfReader

            # 重置文件指针
            file.seek(0)

            # 创建PDF读取器
            pdf_reader = PdfReader(file)

            # 提取所有页面的文本
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"=== 第{page_num + 1}页 ===\n{page_text.strip()}")
                except Exception as e:
                    logger.warning(f"提取PDF第{page_num + 1}页失败: {e}")
                    continue

            content = "\n\n".join(text_content)
            logger.info(f"成功提取PDF内容，页数: {len(pdf_reader.pages)}, 内容长度: {len(content)}")

            return content

        except ImportError:
            logger.error("pypdf库未安装，无法解析PDF文档")
            return ""
        except Exception as e:
            logger.error(f"PDF文档解析失败: {e}")
            # 如果PDF解析失败，不要fallback到文本读取
            return ""

    def _extract_from_word(self, file) -> str:
        """提取Word文件内容，保留标题格式和表格位置"""
        try:
            from docx import Document
            from docx.table import Table
            from docx.text.paragraph import Paragraph

            # 重置文件指针
            file.seek(0)

            # 使用python-docx读取Word文档
            doc = Document(file)

            logger.info(f"开始提取Word文档，段落数: {len(doc.paragraphs)}, 表格数: {len(doc.tables)}")

            # 按文档顺序处理所有元素（段落和表格）
            content_parts = []
            extracted_paragraphs = 0
            extracted_tables = 0

            # 创建元素到对象的映射，避免重复查找（性能优化 + 防止内容丢失）
            paragraph_map = {p._element: p for p in doc.paragraphs}
            table_map = {t._element: t for t in doc.tables}

            # 遍历文档的所有元素
            for element in doc.element.body:
                if element.tag.endswith('p'):  # 段落元素
                    paragraph = paragraph_map.get(element)
                    if paragraph:
                        text = paragraph.text.strip()
                        if text:  # 只处理非空段落
                            markdown_text = self._convert_paragraph_to_markdown(paragraph)
                            content_parts.append(markdown_text)
                            extracted_paragraphs += 1

                elif element.tag.endswith('tbl'):  # 表格元素
                    table = table_map.get(element)
                    if table:
                        table_content = self._extract_table_content(table)
                        if table_content:
                            content_parts.append(table_content)
                            extracted_tables += 1

            content = "\n\n".join(content_parts)

            logger.info(f"Word文档提取完成 - 提取段落: {extracted_paragraphs}, 提取表格: {extracted_tables}, 总内容长度: {len(content)}")
            logger.info(f"原始段落数: {len(doc.paragraphs)}, 实际提取: {extracted_paragraphs}, 差异: {len(doc.paragraphs) - extracted_paragraphs}")

            # 如果内容长度为0或者提取的段落数明显少于原始段落数，记录警告
            if len(content) == 0:
                logger.error("Word文档提取结果为空！")
            elif extracted_paragraphs < len(doc.paragraphs) * 0.5:
                logger.warning(f"Word文档可能存在内容丢失！原始{len(doc.paragraphs)}段，仅提取{extracted_paragraphs}段")

            return content

        except ImportError:
            logger.error("python-docx库未安装，无法解析Word文档")
            return ""
        except Exception as e:
            logger.error(f"Word文档解析失败: {e}")
            import traceback
            logger.error(f"详细错误: {traceback.format_exc()}")
            # 如果解析失败，使用简化方法
            return self._extract_from_word_simple(file)

    def _extract_from_word_simple(self, file) -> str:
        """简化的Word文档提取方法（备用）"""
        try:
            from docx import Document

            # 重置文件指针
            file.seek(0)
            doc = Document(file)

            # 简单提取段落
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    markdown_text = self._convert_paragraph_to_markdown(paragraph)
                    paragraphs.append(markdown_text)

            # 提取表格
            for table in doc.tables:
                table_content = self._extract_table_content(table)
                if table_content:
                    paragraphs.append(table_content)

            content = "\n\n".join(paragraphs)
            logger.warning(f"使用简化方法提取Word文档内容，段落数: {len(paragraphs)}")
            return content

        except Exception as e:
            logger.error(f"简化Word文档解析也失败: {e}")
            return ""

    def _convert_paragraph_to_markdown(self, paragraph) -> str:
        """将Word段落转换为Markdown格式"""
        text = paragraph.text.strip()
        if not text:
            return ""

        # 检查段落样式
        style_name = paragraph.style.name if paragraph.style else ""

        # 根据样式名称转换为Markdown标题
        if 'Heading 1' in style_name or 'heading 1' in style_name.lower():
            return f"# {text}"
        elif 'Heading 2' in style_name or 'heading 2' in style_name.lower():
            return f"## {text}"
        elif 'Heading 3' in style_name or 'heading 3' in style_name.lower():
            return f"### {text}"
        elif 'Heading 4' in style_name or 'heading 4' in style_name.lower():
            return f"#### {text}"
        elif 'Heading 5' in style_name or 'heading 5' in style_name.lower():
            return f"##### {text}"
        elif 'Heading 6' in style_name or 'heading 6' in style_name.lower():
            return f"###### {text}"
        else:
            # 对于非标题样式的段落，直接返回原文本
            # 不再根据粗体格式推测标题，避免误判
            return text



    def _extract_table_content(self, table, depth=0) -> str:
        """提取表格内容为Markdown格式，支持嵌套表格
        
        Args:
            table: Word表格对象
            depth: 递归深度，用于处理嵌套表格
        """
        try:
            table_rows = []
            total_cells = 0
            extracted_cells = 0
            nested_tables_found = 0

            for i, row in enumerate(table.rows):
                row_cells = []
                for cell in row.cells:
                    total_cells += 1
                    
                    # 检查单元格中是否包含嵌套表格
                    nested_tables = cell.tables
                    if nested_tables and depth < 3:  # 最多支持3层嵌套
                        nested_tables_found += len(nested_tables)
                        # 提取单元格中的文本（不包括嵌套表格的文本）
                        cell_text_parts = []
                        
                        # 提取单元格段落文本
                        for paragraph in cell.paragraphs:
                            para_text = paragraph.text.strip()
                            if para_text:
                                cell_text_parts.append(para_text)
                        
                        # 递归提取嵌套表格
                        for nested_table in nested_tables:
                            nested_content = self._extract_table_content(nested_table, depth + 1)
                            if nested_content:
                                cell_text_parts.append(f"[嵌套表格]\n{nested_content}")
                        
                        cell_text = " ".join(cell_text_parts).replace('\n', ' ')
                    else:
                        # 普通单元格，直接提取文本
                        cell_text = cell.text.strip().replace('\n', ' ')
                    
                    if cell_text:
                        row_cells.append(cell_text)
                        extracted_cells += 1
                    else:
                        row_cells.append("")

                if any(cell for cell in row_cells):  # 如果行不为空
                    table_rows.append(" | ".join(row_cells))

                    # 如果是第一行，添加分隔符
                    if i == 0 and len(row_cells) > 0:
                        separator = " | ".join(["---"] * len(row_cells))
                        table_rows.append(separator)

            if table_rows:
                table_content = "\n".join(table_rows)
                indent = "  " * depth
                logger.info(f"{indent}表格提取完成 (深度{depth}) - 行数: {len(table.rows)}, 总单元格: {total_cells}, 非空单元格: {extracted_cells}, 嵌套表格: {nested_tables_found}, 内容长度: {len(table_content)}")
                return table_content
            else:
                logger.warning(f"表格为空或所有行都为空 (深度{depth})")
            return ""

        except Exception as e:
            logger.error(f"表格内容提取失败 (深度{depth}): {e}")
            import traceback
            logger.error(f"详细错误: {traceback.format_exc()}")
            return ""

    def _get_sample_content(self) -> str:
        """获取示例内容用于测试"""
        return """
        # 电商平台需求文档

        ## 1. 用户管理模块
        
        ### 1.1 用户注册
        用户可以通过手机号或邮箱注册账户，需要验证码验证。
        注册时需要填写基本信息：用户名、密码、确认密码。
        密码强度要求：至少8位，包含字母和数字。
        
        ### 1.2 用户登录
        支持手机号/邮箱/用户名登录。
        支持记住登录状态，7天内免登录。
        连续登录失败5次后锁定账户30分钟。
        
        ### 1.3 权限管理
        系统分为普通用户、VIP用户、管理员三种角色。
        不同角色拥有不同的功能权限。
        
        ## 2. 商品管理模块
        
        ### 2.1 商品展示
        首页展示热门商品、新品推荐、分类商品。
        商品详情页包含图片、价格、规格、评价等信息。
        支持商品搜索和筛选功能。
        
        ### 2.2 商品分类
        支持多级分类管理。
        分类支持图标和描述。
        
        ## 3. 订单管理模块
        
        ### 3.1 购物车
        用户可以添加商品到购物车。
        支持修改商品数量、删除商品。
        购物车数据在登录状态下同步。
        
        ### 3.2 订单创建
        用户选择商品后可以创建订单。
        需要选择收货地址和支付方式。
        支持优惠券和积分抵扣。
        
        ### 3.3 订单状态管理
        订单状态包括：待支付、已支付、已发货、已收货、已完成、已取消。
        用户可以查看订单详情和物流信息。
        
        ## 4. 支付管理模块
        
        ### 4.1 支付方式
        支持微信支付、支付宝、银行卡支付。
        支持分期付款（信用卡）。
        
        ### 4.2 支付安全
        所有支付信息加密传输。
        支付密码验证。
        异常支付行为监控和风控。
        
        ## 5. 库存管理模块
        
        ### 5.1 库存监控
        实时监控商品库存数量。
        库存不足时自动预警。
        
        ### 5.2 库存更新
        支持手动和自动库存更新。
        订单完成后自动扣减库存。
        """
    
    def preprocess_content(self, content: str) -> str:
        """预处理文档内容"""
        # 清理多余的空白字符
        content = re.sub(r'\n\s*\n', '\n\n', content)
        content = re.sub(r' +', ' ', content)

        # 不再统一标题格式！保持原有的标题层级
        # 这个正则表达式是导致所有标题都变成##的罪魁祸首
        # content = re.sub(r'^#+\s*', '## ', content, flags=re.MULTILINE)

        return content.strip()


class ModuleSplitter:
    """模块拆分器 - 负责AI智能模块识别和拆分"""

    def __init__(self, user=None):
        self.user = user
        self.llm = self._get_llm_instance()
    
    def _get_llm_instance(self):
        """获取LLM实例"""
        try:
            active_config = LLMConfig.objects.filter(is_active=True).first()
            if not active_config:
                raise Exception("没有可用的LLM配置")

            # 使用新的LLM工厂函数，支持多供应商
            return create_llm_instance(active_config, temperature=0.1)
        except Exception as e:
            logger.error(f"获取LLM实例失败: {e}")
            raise

    def _get_user_prompt(self, prompt_type: str) -> str:
        """获取用户的提示词，如果没有则返回None"""
        if self.user:
            user_prompt = UserPrompt.get_user_prompt_by_type(self.user, prompt_type)
            if user_prompt:
                return user_prompt.content

        # 没有找到用户提示词，返回None
        return None


    
    def split_into_modules(self, document: RequirementDocument, content: str,
                          split_options: dict = None) -> List[Dict[str, Any]]:
        """将文档拆分为功能模块 - 支持多种拆分方式"""
        try:
            # 解析拆分选项
            split_options = split_options or {}
            split_level = split_options.get('split_level', 'auto')
            include_context = split_options.get('include_context', True)
            chunk_size = split_options.get('chunk_size', 2000)

            logger.info(f"开始拆分文档，拆分级别: {split_level}, 包含上下文: {include_context}")

            # 根据拆分级别选择方法
            if split_level == 'auto':
                modules_data = self._split_by_character_length(content, chunk_size)
            elif split_level in ['h1', 'h2', 'h3']:
                modules_data = self._split_by_heading_level(content, split_level, include_context)
            else:
                raise ValueError(f"不支持的拆分级别: {split_level}")

            logger.info(f"拆分完成，生成 {len(modules_data)} 个模块")
            return modules_data

        except Exception as e:
            logger.error(f"模块拆分失败: {e}")
            raise
    
    def _analyze_document_structure(self, content: str) -> List[Dict[str, Any]]:
        """分析文档结构，识别模块边界"""

        # 从数据库获取用户的提示词
        structure_prompt = self._get_user_prompt('document_structure')
        if not structure_prompt:
            raise ValueError("用户未配置文档结构分析提示词，请先在提示词管理中配置")
        
        try:
            # 如果内容太长，进行智能截断，但保持结构完整
            processed_content = self._prepare_content_for_analysis(content)

            prompt_template = Template(structure_prompt)
            formatted_prompt = prompt_template.safe_substitute(content=processed_content)
            messages = [
                SystemMessage(content="你是一个专业的需求分析师，擅长分析需求文档结构。"),
                HumanMessage(content=formatted_prompt)
            ]
            
            response = self.llm.invoke(messages)
            
            # 提取JSON内容
            json_match = re.search(r'```json\s*(.*?)\s*```', response.content, re.DOTALL)
            if json_match:
                json_str = json_match.group(1)
            else:
                # 如果没有找到代码块，尝试直接解析整个响应
                json_str = response.content.strip()
            
            modules_structure = json.loads(json_str)
            logger.info(f"识别到 {len(modules_structure)} 个模块")
            
            return modules_structure

        except json.JSONDecodeError as e:
            logger.error(f"解析模块结构JSON失败: {e}")
            # 返回默认的模块结构
            return self._get_default_modules_structure(content)
        except Exception as e:
            logger.error(f"分析文档结构失败: {e}")
            return self._get_default_modules_structure(content)

    def _prepare_content_for_analysis(self, content: str) -> str:
        """为AI分析准备内容，确保结构完整"""
        # 如果内容不长，直接返回
        if len(content) <= 6000:
            return content

        # 如果内容太长，按标题进行智能截断
        lines = content.split('\n')
        processed_lines = []
        current_length = 0

        for line in lines:
            # 保留所有标题行
            if line.strip().startswith('#') or line.strip().startswith('##'):
                processed_lines.append(line)
                current_length += len(line) + 1
            # 保留重要内容，但限制总长度
            elif current_length < 5500:
                processed_lines.append(line)
                current_length += len(line) + 1
            # 如果遇到新的标题，即使超长也要保留
            elif line.strip().startswith('#'):
                processed_lines.append(line)
                processed_lines.append("... (内容已截断，但保留了所有标题结构)")
                break

        return '\n'.join(processed_lines)
    
    def _extract_module_content(self, content: str, modules_structure: List[Dict]) -> List[Dict[str, Any]]:
        """根据模块结构提取具体内容"""
        modules_data = []
        
        for i, module_info in enumerate(modules_structure):
            try:
                # 根据标识符提取模块内容
                module_content = self._extract_content_by_markers(
                    content, 
                    module_info.get('start_marker', ''),
                    module_info.get('end_marker', ''),
                    i, 
                    len(modules_structure)
                )
                
                if module_content:
                    modules_data.append({
                        'title': module_info['title'],
                        'content': module_content,
                        'description': module_info.get('description', ''),
                        'order': i + 1,
                        'confidence_score': module_info.get('confidence', 0.8),
                        'estimated_complexity': module_info.get('estimated_complexity', 'medium'),
                        'is_auto_generated': True
                    })
                    
            except Exception as e:
                logger.error(f"提取模块内容失败: {e}")
                continue
        
        return modules_data
    
    def _extract_content_by_markers(self, content: str, start_marker: str, end_marker: str,
                                   index: int, total_modules: int) -> str:
        """根据标识符提取内容 - 使用简单的基于标题的拆分方法"""
        try:
            # 使用简单的基于标题结构的拆分，类似知识库的做法
            return self._extract_content_by_simple_structure(content, index, total_modules)

        except Exception as e:
            logger.error(f"根据标识符提取内容失败: {e}")
            return ""



    def _extract_content_by_title_structure(self, content: str, start_marker: str,
                                          index: int, total_modules: int) -> str:
        """基于标题结构的内容提取方法"""
        lines = content.split('\n')

        # 找到所有主要标题（## 级别，但不包括 ###）
        title_positions = []
        title_texts = []

        for i, line in enumerate(lines):
            stripped_line = line.strip()
            if stripped_line.startswith('##') and not stripped_line.startswith('###'):
                title_positions.append(i)
                title_texts.append(stripped_line)

        logger.info(f"找到 {len(title_positions)} 个主要标题: {title_texts}")

        if not title_positions:
            # 如果没有找到标题，平均分割
            total_lines = len(lines)
            lines_per_module = total_lines // total_modules
            start_idx = index * lines_per_module
            end_idx = (index + 1) * lines_per_module if index < total_modules - 1 else total_lines
            extracted = '\n'.join(lines[start_idx:end_idx]).strip()
            logger.info(f"模块 {index}: 平均分割，行数 {start_idx}-{end_idx}")
            return extracted

        # 根据标题位置提取内容
        if index < len(title_positions):
            start_idx = title_positions[index]
            end_idx = title_positions[index + 1] if index + 1 < len(title_positions) else len(lines)
            extracted = '\n'.join(lines[start_idx:end_idx]).strip()

            logger.info(f"模块 {index}: 标题 '{title_texts[index]}', 行数 {start_idx}-{end_idx}, 内容长度 {len(extracted)}")
            return extracted

        logger.warning(f"模块 {index} 超出标题数量范围")
        return ""

    def _extract_content_by_simple_structure(self, content: str, index: int, total_modules: int) -> str:
        """简单的基于结构的内容提取 - 类似知识库的纯拆分方法"""
        lines = content.split('\n')

        # 找到所有可能的模块分界点（主要标题）
        section_boundaries = []

        for i, line in enumerate(lines):
            stripped_line = line.strip()

            # 识别主要的功能模块标题
            if self._is_main_section_title(stripped_line):
                section_boundaries.append(i)

        logger.info(f"找到 {len(section_boundaries)} 个主要章节边界")

        # 如果没有找到明显的章节，按内容长度平均分割
        if len(section_boundaries) < 2:
            return self._split_by_content_length(content, index, total_modules)

        # 根据章节边界提取内容
        if index < len(section_boundaries):
            start_idx = section_boundaries[index]
            end_idx = section_boundaries[index + 1] if index + 1 < len(section_boundaries) else len(lines)

            # 提取原始内容，不做任何修改
            extracted_lines = lines[start_idx:end_idx]
            extracted_content = '\n'.join(extracted_lines)

            logger.info(f"模块 {index}: 章节拆分，行数 {start_idx}-{end_idx}, 内容长度 {len(extracted_content)}")
            return extracted_content

        # 如果索引超出范围，返回空
        return ""

    def _is_main_section_title(self, line: str) -> bool:
        """判断是否为主要章节标题"""
        if not line:
            return False

        # 常见的功能模块关键词
        module_keywords = [
            '管理', '模块', '系统', '功能', '服务', '平台', '中心',
            '注册', '登录', '用户', '数据', '统计', '报表', '分析',
            '权限', '角色', '机构', '菜单', '日志', '配置', '设置'
        ]

        # 检查是否包含模块关键词且长度适中
        if any(keyword in line for keyword in module_keywords) and len(line) < 50:
            return True

        # 检查是否为数字编号的标题（如 "1. 用户管理"）
        if line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            return True

        # 检查是否为Markdown标题
        if line.startswith('##') and not line.startswith('###'):
            return True

        return False

    def _split_by_content_length(self, content: str, index: int, total_modules: int) -> str:
        """按内容长度平均分割 - 保持原文不变"""
        lines = content.split('\n')
        total_lines = len(lines)

        # 计算每个模块的大致行数
        lines_per_module = total_lines // total_modules

        start_idx = index * lines_per_module
        end_idx = (index + 1) * lines_per_module if index < total_modules - 1 else total_lines

        # 尝试在段落边界调整切分点
        if end_idx < total_lines:
            # 向后查找空行作为更好的切分点
            for i in range(end_idx, min(end_idx + 10, total_lines)):
                if not lines[i].strip():
                    end_idx = i
                    break

        extracted_lines = lines[start_idx:end_idx]
        extracted_content = '\n'.join(extracted_lines)

        logger.info(f"模块 {index}: 长度拆分，行数 {start_idx}-{end_idx}, 内容长度 {len(extracted_content)}")
        return extracted_content

    def _split_by_document_structure(self, content: str) -> List[Dict[str, Any]]:
        """基于文档结构的纯文本拆分 - 严格按文档顺序"""
        lines = content.split('\n')

        # 找到所有二级标题（## 开头的）作为主要模块分界点
        module_boundaries = []
        module_titles = []

        for i, line in enumerate(lines):
            stripped_line = line.strip()

            # 只识别二级标题作为模块边界，确保按文档顺序
            if stripped_line.startswith('## ') and not stripped_line.startswith('### '):
                module_boundaries.append(i)
                module_titles.append(stripped_line)

        logger.info(f"识别到 {len(module_boundaries)} 个二级标题: {module_titles}")

        # 如果没有找到二级标题，按其他方式分割
        if len(module_boundaries) < 2:
            return self._split_by_content_sections(content)

        # 按模块边界拆分内容，严格按顺序
        modules_data = []

        # 第一个模块：从文档开始到第一个二级标题
        if module_boundaries[0] > 0:
            first_module_lines = lines[0:module_boundaries[0]]
            first_module_content = '\n'.join(first_module_lines).strip()

            if first_module_content:
                modules_data.append({
                    'title': '文档概述',
                    'content': first_module_content,
                    'confidence_score': 0.90,
                    'estimated_complexity': 'low',
                    'order': 1,
                    'start_page': 1,
                    'end_page': 1
                })
                logger.info(f"模块 1: 文档概述, 内容长度: {len(first_module_content)}")

        # 其他模块：按二级标题分割
        for i, start_idx in enumerate(module_boundaries):
            end_idx = module_boundaries[i + 1] if i + 1 < len(module_boundaries) else len(lines)

            # 提取模块内容，保持原文不变
            module_lines = lines[start_idx:end_idx]
            module_content = '\n'.join(module_lines).strip()

            # 生成模块标题（去掉##前缀）
            title = module_titles[i].replace('## ', '').strip()

            modules_data.append({
                'title': title,
                'content': module_content,
                'confidence_score': 0.95,
                'estimated_complexity': 'medium',
                'order': len(modules_data) + 1,
                'start_page': 1,
                'end_page': 1
            })

            logger.info(f"模块 {len(modules_data)}: {title}, 内容长度: {len(module_content)}")

        return modules_data

    def _is_functional_module_title(self, line: str) -> bool:
        """判断是否为功能模块标题"""
        if not line or len(line) > 100:
            return False

        # 功能模块的关键词
        functional_keywords = [
            '用户', '登录', '注册', '管理', '系统', '数据', '统计', '报表',
            '权限', '角色', '机构', '菜单', '日志', '配置', '设置', '服务',
            '平台', '中心', '模块', '功能', '业务', '流程', '接口', 'API'
        ]

        # 检查是否包含功能关键词
        has_functional_keyword = any(keyword in line for keyword in functional_keywords)

        # 检查是否为合理的标题格式
        is_title_format = (
            line.endswith('管理') or
            line.endswith('模块') or
            line.endswith('系统') or
            line.endswith('功能') or
            line.endswith('服务') or
            '管理' in line or
            '模块' in line
        )

        return has_functional_keyword and is_title_format

    def _split_by_content_sections(self, content: str) -> List[Dict[str, Any]]:
        """按内容章节智能分割"""
        lines = content.split('\n')

        # 找到所有可能的章节分界点
        section_points = []

        for i, line in enumerate(lines):
            stripped_line = line.strip()

            # 识别章节标题（包括数字编号、标题等）
            if (stripped_line and
                (stripped_line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')) or
                 stripped_line.startswith('##') or
                 (len(stripped_line) < 50 and not stripped_line.startswith(('-', '•', '*'))))):
                section_points.append(i)

        logger.info(f"找到 {len(section_points)} 个章节分界点")

        # 如果章节太少，按内容长度分割
        if len(section_points) < 3:
            return self._split_by_equal_parts(content, 5)  # 默认分成5个模块

        # 按章节分割
        modules_data = []
        for i, start_idx in enumerate(section_points):
            end_idx = section_points[i + 1] if i + 1 < len(section_points) else len(lines)

            module_lines = lines[start_idx:end_idx]
            module_content = '\n'.join(module_lines)

            # 生成标题
            first_line = lines[start_idx].strip()
            title = first_line if len(first_line) < 50 else f"章节{i+1}"

            modules_data.append({
                'title': title,
                'content': module_content,
                'confidence_score': 0.85,
                'estimated_complexity': 'medium',
                'order': i + 1,
                'start_page': 1,
                'end_page': 1
            })

        return modules_data

    def _split_by_character_length(self, content: str, chunk_size: int = 2000) -> List[Dict[str, Any]]:
        """纯粹按字数拆分，保证内容一个字都不改"""
        modules_data = []
        start = 0
        order = 1

        logger.info(f"开始按字数拆分文档，总长度: {len(content)}, 分块大小: {chunk_size}")

        while start < len(content):
            end = min(start + chunk_size, len(content))

            # 尝试在合适的边界切分（段落、句子边界）
            if end < len(content):
                # 向后查找最近的段落边界（双换行）
                for i in range(end, min(end + 200, len(content))):
                    if i + 1 < len(content) and content[i:i+2] == '\n\n':
                        end = i + 1
                        break
                else:
                    # 如果没找到段落边界，查找单换行
                    for i in range(end, min(end + 100, len(content))):
                        if content[i] == '\n':
                            end = i + 1
                            break
                    else:
                        # 如果没找到换行，查找句号
                        for i in range(end, min(end + 50, len(content))):
                            if content[i] in '。！？.!?':
                                end = i + 1
                                break

            # 提取分块内容，保持原文完全不变
            chunk_content = content[start:end]

            # 计算页数（假设每页约500字符）
            start_page = (start // 500) + 1
            end_page = (end // 500) + 1

            # 生成模块数据
            modules_data.append({
                'title': f'第{order}部分',
                'content': chunk_content,
                'confidence_score': 1.0,  # 按字数拆分，置信度100%
                'estimated_complexity': 'medium',
                'order': order,
                'start_page': start_page,
                'end_page': end_page,
                'start_position': start,  # 记录字符位置，方便用户调整
                'end_position': end
            })

            logger.info(f"分块 {order}: 字符位置 {start}-{end}, 长度: {len(chunk_content)}")

            start = end
            order += 1

        logger.info(f"拆分完成，共生成 {len(modules_data)} 个分块")
        return modules_data

    def _split_by_heading_level(self, content: str, level: str, include_context: bool = True) -> List[Dict[str, Any]]:
        """根据标题级别拆分文档"""
        lines = content.split('\n')

        # 标题模式映射
        heading_patterns = {
            'h1': '# ',
            'h2': '## ',
            'h3': '### '
        }

        target_pattern = heading_patterns.get(level)
        if not target_pattern:
            raise ValueError(f"不支持的标题级别: {level}")

        logger.info(f"按 {level.upper()} 级别标题拆分文档")

        # 查找所有目标级别的标题
        heading_positions = []
        context_headings = {}  # 存储上级标题作为上下文

        current_h1 = ""
        current_h2 = ""

        for i, line in enumerate(lines):
            stripped = line.strip()

            # 跟踪上级标题
            if stripped.startswith('# ') and not stripped.startswith('## '):
                current_h1 = stripped
            elif stripped.startswith('## ') and not stripped.startswith('### '):
                current_h2 = stripped

            # 找到目标级别的标题
            if stripped.startswith(target_pattern) and not stripped.startswith(target_pattern + '#'):
                heading_positions.append(i)

                # 保存上下文
                context = []
                if include_context:
                    if level == 'h2' and current_h1:
                        context.append(current_h1)
                    elif level == 'h3' and current_h1:
                        context.append(current_h1)
                        if current_h2:
                            context.append(current_h2)

                context_headings[i] = context

        logger.info(f"找到 {len(heading_positions)} 个 {level.upper()} 级别标题")

        if len(heading_positions) == 0:
            # 没有找到目标级别的标题，fallback到字数拆分
            logger.warning(f"未找到 {level.upper()} 级别标题，使用字数拆分")
            return self._split_by_character_length(content)

        # 按标题拆分内容
        modules_data = []
        order = 1

        # 处理第一个目标标题之前的内容（前言部分）
        if heading_positions and heading_positions[0] > 0:
            preface_lines = lines[:heading_positions[0]]
            preface_content = '\n'.join(preface_lines).strip()

            if preface_content:  # 如果前言部分有内容
                # 尝试从前言中提取标题
                preface_title = "前言"
                for line in preface_lines:
                    stripped = line.strip()
                    if stripped.startswith('# ') and not stripped.startswith('## '):
                        preface_title = stripped.replace('# ', '').strip()
                        break

                start_char = 0
                end_char = len('\n'.join(lines[:heading_positions[0]]))
                start_page = 1
                end_page = (end_char // 500) + 1

                modules_data.append({
                    'title': preface_title,
                    'content': preface_content,
                    'confidence_score': 0.90,
                    'estimated_complexity': 'medium',
                    'order': order,
                    'start_page': start_page,
                    'end_page': end_page,
                    'start_position': start_char,
                    'end_position': end_char,
                    'split_method': f'{level}_heading_preface'
                })

                logger.info(f"模块 {order}: {preface_title} (前言), 内容长度: {len(preface_content)}")
                order += 1

        # 处理每个目标级别标题的内容
        for i, start_idx in enumerate(heading_positions):
            end_idx = heading_positions[i + 1] if i + 1 < len(heading_positions) else len(lines)

            # 构建模块内容（从目标标题开始到下一个同级标题）
            module_lines = lines[start_idx:end_idx]
            module_content = '\n'.join(module_lines).strip()

            # 生成模块标题
            title_line = lines[start_idx].strip()
            title = title_line.replace(target_pattern, '').strip()

            # 计算字符位置和页数
            start_char = len('\n'.join(lines[:start_idx]))
            end_char = len('\n'.join(lines[:end_idx]))
            start_page = (start_char // 500) + 1
            end_page = (end_char // 500) + 1

            modules_data.append({
                'title': title,
                'content': module_content,
                'confidence_score': 0.95,  # 按标题拆分，置信度很高
                'estimated_complexity': 'medium',
                'order': order,
                'start_page': start_page,
                'end_page': end_page,
                'start_position': start_char,
                'end_position': end_char,
                'split_method': f'{level}_heading'
            })

            logger.info(f"模块 {order}: {title}, 内容长度: {len(module_content)}")
            order += 1

        return modules_data

    def _split_by_equal_parts(self, content: str, num_parts: int) -> List[Dict[str, Any]]:
        """按相等部分分割内容"""
        lines = content.split('\n')
        total_lines = len(lines)
        lines_per_part = total_lines // num_parts

        modules_data = []
        for i in range(num_parts):
            start_idx = i * lines_per_part
            end_idx = (i + 1) * lines_per_part if i < num_parts - 1 else total_lines

            # 尝试在段落边界调整
            if end_idx < total_lines:
                for j in range(end_idx, min(end_idx + 5, total_lines)):
                    if not lines[j].strip():
                        end_idx = j
                        break

            module_lines = lines[start_idx:end_idx]
            module_content = '\n'.join(module_lines)

            modules_data.append({
                'title': f"部分{i+1}",
                'content': module_content,
                'confidence_score': 0.75,
                'estimated_complexity': 'medium',
                'order': i + 1,
                'start_page': 1,
                'end_page': 1
            })

        return modules_data
    
    def _fuzzy_find_marker(self, content: str, marker: str) -> int:
        """模糊查找标识符"""
        # 简单的模糊匹配逻辑
        marker_words = marker.split()
        if marker_words:
            return content.find(marker_words[0])
        return -1
    
    def _optimize_modules(self, modules_data: List[Dict]) -> List[Dict[str, Any]]:
        """优化模块拆分结果"""
        optimized = []
        
        for module in modules_data:
            # 检查模块内容长度
            content_length = len(module['content'])
            
            if content_length < 100:
                # 内容过短，可能需要合并
                module['confidence_score'] *= 0.7
                module['suggestion'] = '内容较少，建议考虑与其他模块合并'
            elif content_length > 5000:
                # 内容过长，可能需要拆分
                module['confidence_score'] *= 0.8
                module['suggestion'] = '内容较多，建议考虑进一步拆分'
            
            # 计算页码信息（模拟）
            module['start_page'] = module['order']
            module['end_page'] = module['order']
            
            optimized.append(module)
        
        return optimized
    
    def _get_default_modules_structure(self, content: str) -> List[Dict]:
        """获取默认的模块结构（当AI分析失败时使用）"""
        # 基于标题结构的简单拆分
        lines = content.split('\n')
        modules = []
        current_module = None
        
        for line in lines:
            line = line.strip()
            if line.startswith('##') and not line.startswith('###'):
                if current_module:
                    modules.append(current_module)
                
                title = line.replace('##', '').strip()
                current_module = {
                    'title': title,
                    'description': f'{title}相关功能需求',
                    'start_marker': line,
                    'end_marker': '',
                    'confidence_score': 0.6,
                    'estimated_complexity': 'medium',
                    'order': 1
                }
        
        if current_module:
            modules.append(current_module)
        
        return modules if modules else [{
            'title': '需求文档',
            'description': '完整需求文档内容',
            'start_marker': '',
            'end_marker': '',
            'confidence_score': 0.5,
            'estimated_complexity': 'high',
            'order': 1
        }]


class RequirementModuleService:
    """需求模块服务 - 统一的模块管理服务"""

    def __init__(self, user=None):
        self.user = user
        self.document_processor = DocumentProcessor()
        self.module_splitter = ModuleSplitter(user=user)
    
    def process_document_and_split(self, document: RequirementDocument,
                                  split_options: dict = None) -> List[RequirementModule]:
        """处理文档并进行模块拆分"""
        try:
            # 更新文档状态
            document.status = 'module_split'
            document.save()
            
            # 提取文档内容
            logger.info(f"开始处理文档 {document.id}: {document.title}")
            content = self.document_processor.extract_content(document)
            if not content:
                raise Exception("无法提取文档内容")
            
            logger.info(f"文档内容提取完成，原始长度: {len(content)} 字符")
            
            # 预处理内容
            processed_content = self.document_processor.preprocess_content(content)
            logger.info(f"内容预处理完成，处理后长度: {len(processed_content)} 字符")
            
            # 更新文档统计信息
            document.word_count = len(processed_content)
            document.page_count = max(1, (len(processed_content) // 500) + 1)  # 假设每页约500字符
            document.content = processed_content
            document.save()
            
            # 验证保存后的内容
            document.refresh_from_db()
            saved_content_length = len(document.content) if document.content else 0
            logger.info(f"文档保存到数据库后，内容长度: {saved_content_length} 字符")
            
            if saved_content_length < len(processed_content):
                logger.error(f"!!!内容截断警告!!! 处理前: {len(processed_content)}, 保存后: {saved_content_length}, 丢失: {len(processed_content) - saved_content_length} 字符")
            elif saved_content_length == 0:
                logger.error(f"!!!严重错误!!! 文档内容保存后为空")
            
            # 模块拆分（支持多种拆分方式）
            logger.info(f"开始模块拆分，拆分选项: {split_options}")
            modules_data = self.module_splitter.split_into_modules(document, processed_content, split_options)
            
            # 创建模块对象
            modules = []
            for i, module_data in enumerate(modules_data):
                module_content_length = len(module_data['content'])
                logger.info(f"创建模块 {i+1}/{len(modules_data)}: {module_data['title']}, 内容长度: {module_content_length}")
                
                module = RequirementModule.objects.create(
                    document=document,
                    title=module_data['title'],
                    content=module_data['content'],
                    start_page=module_data.get('start_page', 1),
                    end_page=module_data.get('end_page', 1),
                    start_position=module_data.get('start_position'),
                    end_position=module_data.get('end_position'),
                    order=module_data['order'],
                    confidence_score=module_data['confidence_score'],
                    is_auto_generated=True,
                    ai_suggested_title=module_data['title']
                )
                
                # 验证模块保存后的内容
                module.refresh_from_db()
                saved_module_content_length = len(module.content) if module.content else 0
                logger.info(f"模块 {module.id} 保存后，内容长度: {saved_module_content_length}")
                
                if saved_module_content_length < module_content_length:
                    logger.error(f"!!!模块内容截断!!! 模块ID: {module.id}, 原始: {module_content_length}, 保存: {saved_module_content_length}, 丢失: {module_content_length - saved_module_content_length} 字符")
                
                modules.append(module)
            
            # 更新文档状态
            document.status = 'user_reviewing'
            document.save()
            
            logger.info(f"文档 {document.id} 模块拆分完成，生成 {len(modules)} 个模块")
            return modules
            
        except Exception as e:
            logger.error(f"文档模块拆分失败: {e}")
            import traceback
            logger.error(f"详细错误: {traceback.format_exc()}")
            document.status = 'failed'
            document.save()
            raise


class ModuleOperationService:
    """模块操作服务 - 处理用户的模块编辑操作"""

    def __init__(self, document: RequirementDocument):
        self.document = document

    def execute_operation(self, operation_data: dict) -> dict:
        """执行单个模块操作"""
        operation = operation_data['operation']

        if operation == 'merge':
            return self._merge_modules(operation_data)
        elif operation == 'split':
            return self._split_module(operation_data)
        elif operation == 'reorder':
            return self._reorder_modules(operation_data)
        elif operation == 'rename':
            return self._rename_module(operation_data)
        elif operation == 'delete':
            return self._delete_module(operation_data)
        elif operation == 'create':
            return self._create_module(operation_data)
        else:
            raise ValueError(f"不支持的操作类型: {operation}")

    def execute_batch_operations(self, operations_data: List[dict]) -> dict:
        """执行批量模块操作"""
        results = []

        try:
            for operation_data in operations_data:
                result = self.execute_operation(operation_data)
                results.append(result)

            # 重新计算所有模块的排序
            self._normalize_module_orders()

            return {
                'success': True,
                'operations_count': len(results),
                'results': results,
                'message': f'成功执行 {len(results)} 个操作'
            }

        except Exception as e:
            logger.error(f"批量操作失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'completed_operations': len(results)
            }

    def _merge_modules(self, operation_data: dict) -> dict:
        """合并模块"""
        target_module_ids = operation_data['target_modules']
        merge_title = operation_data['merge_title']
        merge_order = operation_data.get('merge_order')

        # 获取目标模块
        modules = RequirementModule.objects.filter(
            id__in=target_module_ids,
            document=self.document
        ).order_by('order')

        if modules.count() < 2:
            raise ValueError("合并操作至少需要2个模块")

        # 合并内容
        merged_content = '\n\n'.join([module.content for module in modules])
        first_module = modules.first()
        last_module = modules.last()

        # 创建新的合并模块
        merged_module = RequirementModule.objects.create(
            document=self.document,
            title=merge_title,
            content=merged_content,
            start_page=first_module.start_page,
            end_page=last_module.end_page,
            order=merge_order or first_module.order,
            is_auto_generated=False,
            confidence_score=0.9  # 用户手动合并，置信度较高
        )

        # 删除原模块
        modules.delete()

        return {
            'operation': 'merge',
            'new_module_id': str(merged_module.id),
            'merged_modules_count': modules.count(),
            'message': f'成功合并 {modules.count()} 个模块为 "{merge_title}"'
        }

    def _split_module(self, operation_data: dict) -> dict:
        """拆分模块"""
        target_module_id = operation_data['target_modules'][0]
        split_points = operation_data['split_points']
        split_titles = operation_data['split_titles']

        # 获取目标模块
        module = RequirementModule.objects.get(
            id=target_module_id,
            document=self.document
        )

        content = module.content

        # 按拆分点分割内容
        content_parts = []
        last_pos = 0

        for split_point in sorted(split_points):
            if split_point > last_pos and split_point < len(content):
                content_parts.append(content[last_pos:split_point])
                last_pos = split_point

        # 添加最后一部分
        content_parts.append(content[last_pos:])

        # 过滤空内容
        content_parts = [part.strip() for part in content_parts if part.strip()]

        if len(content_parts) != len(split_titles):
            raise ValueError("拆分内容数量与标题数量不匹配")

        # 创建新模块
        new_modules = []
        base_order = module.order

        for i, (content_part, title) in enumerate(zip(content_parts, split_titles)):
            new_module = RequirementModule.objects.create(
                document=self.document,
                title=title,
                content=content_part,
                start_page=module.start_page,
                end_page=module.end_page,
                order=base_order + i * 0.1,  # 使用小数保持顺序
                is_auto_generated=False,
                confidence_score=0.8
            )
            new_modules.append(new_module)

        # 删除原模块
        module.delete()

        return {
            'operation': 'split',
            'new_modules': [str(m.id) for m in new_modules],
            'split_count': len(new_modules),
            'message': f'成功将模块拆分为 {len(new_modules)} 个子模块'
        }

    def _reorder_modules(self, operation_data: dict) -> dict:
        """重新排序模块"""
        new_orders = operation_data['new_orders']

        updated_count = 0
        for module_id, new_order in new_orders.items():
            try:
                module = RequirementModule.objects.get(
                    id=module_id,
                    document=self.document
                )
                module.order = new_order
                module.save()
                updated_count += 1
            except RequirementModule.DoesNotExist:
                logger.warning(f"模块 {module_id} 不存在")

        return {
            'operation': 'reorder',
            'updated_count': updated_count,
            'message': f'成功重新排序 {updated_count} 个模块'
        }

    def _rename_module(self, operation_data: dict) -> dict:
        """重命名模块"""
        target_module_id = operation_data['target_modules'][0]
        new_title = operation_data['new_module_data']['title']

        module = RequirementModule.objects.get(
            id=target_module_id,
            document=self.document
        )

        old_title = module.title
        module.title = new_title
        module.is_auto_generated = False
        module.save()

        return {
            'operation': 'rename',
            'module_id': str(module.id),
            'old_title': old_title,
            'new_title': new_title,
            'message': f'成功将模块从 "{old_title}" 重命名为 "{new_title}"'
        }

    def _delete_module(self, operation_data: dict) -> dict:
        """删除模块"""
        target_module_id = operation_data['target_modules'][0]

        module = RequirementModule.objects.get(
            id=target_module_id,
            document=self.document
        )

        title = module.title
        module.delete()

        return {
            'operation': 'delete',
            'deleted_module_id': str(target_module_id),
            'deleted_title': title,
            'message': f'成功删除模块 "{title}"'
        }

    def _create_module(self, operation_data: dict) -> dict:
        """创建新模块"""
        module_data = operation_data['new_module_data']

        module = RequirementModule.objects.create(
            document=self.document,
            title=module_data['title'],
            content=module_data.get('content', ''),
            start_page=module_data.get('start_page', 1),
            end_page=module_data.get('end_page', 1),
            order=module_data.get('order', self._get_next_order()),
            is_auto_generated=False,
            confidence_score=1.0  # 用户手动创建，置信度最高
        )

        return {
            'operation': 'create',
            'new_module_id': str(module.id),
            'title': module.title,
            'message': f'成功创建新模块 "{module.title}"'
        }

    def _get_next_order(self) -> int:
        """获取下一个排序号"""
        last_module = self.document.modules.order_by('-order').first()
        return (last_module.order + 1) if last_module else 1

    def _normalize_module_orders(self):
        """规范化模块排序（确保连续的整数）"""
        modules = self.document.modules.order_by('order')
        for i, module in enumerate(modules, 1):
            if module.order != i:
                module.order = i
                module.save()


class RequirementReviewEngine:
    """需求评审AI分析引擎 - 专业的需求文档评审分析"""

    def __init__(self, user=None):
        self.user = user
        self.llm = self._get_llm_instance()

    def _get_llm_instance(self):
        """获取LLM实例"""
        try:
            active_config = LLMConfig.objects.filter(is_active=True).first()
            if not active_config:
                raise Exception("没有可用的LLM配置")

            # 使用新的LLM工厂函数，支持多供应商
            return create_llm_instance(active_config, temperature=0.1)
        except Exception as e:
            logger.error(f"获取LLM实例失败: {e}")
            raise

    def _get_user_prompt(self, prompt_type: str) -> str:
        """获取用户的提示词，如果没有则返回None"""
        if self.user:
            user_prompt = UserPrompt.get_user_prompt_by_type(self.user, prompt_type)
            if user_prompt:
                return user_prompt.content

        # 没有找到用户提示词，返回None
        return None



    def analyze_document_directly(self, content: str, analysis_options: dict = None) -> dict:
        """直接分析整个文档（不拆分模块）"""
        try:
            # 从数据库获取用户的提示词
            direct_prompt = self._get_user_prompt('direct_analysis')
            if not direct_prompt:
                raise ValueError("用户未配置直接分析提示词，请先在提示词管理中配置")

            prompt_template = Template(direct_prompt)
            formatted_prompt = prompt_template.safe_substitute(content=content[:8000])
            messages = [
                SystemMessage(content="你是一位专业的需求分析师，擅长需求文档评审。"),
                HumanMessage(content=formatted_prompt)
            ]

            response = self.llm.invoke(messages)

            # 提取JSON内容
            json_match = re.search(r'```json\s*(.*?)\s*```', response.content, re.DOTALL)
            if json_match:
                json_str = json_match.group(1)
                analysis_result = json.loads(json_str)

                # 确保必要字段存在
                analysis_result.setdefault('overall_rating', 'average')
                analysis_result.setdefault('completion_score', 70)
                analysis_result.setdefault('clarity_score', 70)
                analysis_result.setdefault('consistency_score', 70)
                analysis_result.setdefault('completeness_score', 70)
                analysis_result.setdefault('summary', '文档评审完成')
                analysis_result.setdefault('recommendations', '请根据问题列表进行改进')
                analysis_result.setdefault('issues', [])

                return analysis_result
            else:
                # 如果没有找到JSON，返回默认结果
                return self._get_default_direct_analysis()

        except json.JSONDecodeError as e:
            logger.error(f"解析直接分析JSON失败: {e}")
            return self._get_default_direct_analysis()
        except Exception as e:
            logger.error(f"直接文档分析失败: {e}")
            return self._get_default_direct_analysis()

    def _get_default_direct_analysis(self) -> dict:
        """获取默认的直接分析结果"""
        return {
            'overall_rating': 'average',
            'completion_score': 70,
            'clarity_score': 70,
            'consistency_score': 70,
            'completeness_score': 70,
            'summary': '文档评审完成，整体质量中等',
            'recommendations': '建议进一步完善需求描述的清晰度和完整性',
            'issues': [
                {
                    'title': '评审系统异常',
                    'description': '自动评审过程中出现异常，建议人工复核',
                    'priority': 'medium',
                    'category': 'system',
                    'location': '整个文档',
                    'suggestion': '请联系管理员检查系统状态'
                }
            ]
        }

    def analyze_completeness(self, content: str) -> dict:
        """完整性专项分析 - 分析完整文档的完整性"""
        logger.info("开始执行完整性分析...")
        completeness_prompt = self._get_user_prompt('completeness_analysis')
        if not completeness_prompt:
            logger.warning("用户未配置完整性分析提示词，返回默认结果")
            return self._get_default_analysis_result('completeness_analysis')
        
        try:
            prompt_template = Template(completeness_prompt)
            formatted_prompt = prompt_template.safe_substitute(document=content)
            logger.debug(f"完整性分析提示词已格式化，文档长度: {len(content)}")
            
            messages = [
                SystemMessage(content="你是一位资深的需求分析专家。"),
                HumanMessage(content=formatted_prompt)
            ]
            
            logger.info("调用LLM进行完整性分析...")
            response = self.llm.invoke(messages)
            logger.info(f"LLM响应完成，内容长度: {len(response.content)}")
            
            json_match = re.search(r'```json\s*(.*?)\s*```', response.content, re.DOTALL)
            
            if json_match:
                result = json.loads(json_match.group(1))
                logger.info(f"完整性分析完成，评分: {result.get('overall_score', 'N/A')}, 问题数: {len(result.get('issues', []))}")
                return result
            else:
                logger.warning("完整性分析响应中未找到JSON格式，使用默认结果")
                logger.debug(f"AI响应内容前500字符: {response.content[:500]}")
                return self._get_default_analysis_result('completeness_analysis')
                
        except Exception as e:
            logger.error(f"完整性分析失败: {e}")
            import traceback
            logger.error(f"详细错误: {traceback.format_exc()}")
            return self._get_default_analysis_result('completeness_analysis')
    
    def analyze_consistency(self, content: str) -> dict:
        """一致性专项分析 - 分析完整文档的一致性"""
        logger.info("开始执行一致性分析...")
        consistency_prompt = self._get_user_prompt('consistency_analysis')
        if not consistency_prompt:
            logger.warning("用户未配置一致性分析提示词，返回默认结果")
            return self._get_default_analysis_result('consistency_analysis')
        
        try:
            prompt_template = Template(consistency_prompt)
            formatted_prompt = prompt_template.safe_substitute(document=content)
            logger.debug(f"一致性分析提示词已格式化，文档长度: {len(content)}")
            
            messages = [
                SystemMessage(content="你是一位资深的需求一致性分析专家。"),
                HumanMessage(content=formatted_prompt)
            ]
            
            logger.info("调用LLM进行一致性分析...")
            response = self.llm.invoke(messages)
            logger.info(f"LLM响应完成，内容长度: {len(response.content)}")
            
            json_match = re.search(r'```json\s*(.*?)\s*```', response.content, re.DOTALL)
            
            if json_match:
                result = json.loads(json_match.group(1))
                logger.info(f"一致性分析完成，评分: {result.get('overall_score', 'N/A')}, 问题数: {len(result.get('issues', []))}")
                return result
            else:
                logger.warning("一致性分析响应中未找到JSON格式，使用默认结果")
                logger.debug(f"AI响应内容前500字符: {response.content[:500]}")
                return self._get_default_analysis_result('consistency_analysis')
                
        except Exception as e:
            logger.error(f"一致性分析失败: {e}")
            import traceback
            logger.error(f"详细错误: {traceback.format_exc()}")
            return self._get_default_analysis_result('consistency_analysis')
    
    def analyze_testability(self, content: str) -> dict:
        """可测性专项分析 - 分析完整文档的可测试性"""
        logger.info("开始执行可测性分析...")
        testability_prompt = self._get_user_prompt('testability_analysis')
        if not testability_prompt:
            logger.warning("用户未配置可测性分析提示词，返回默认结果")
            return self._get_default_analysis_result('testability_analysis')
        
        try:
            prompt_template = Template(testability_prompt)
            formatted_prompt = prompt_template.safe_substitute(document=content)
            logger.debug(f"可测性分析提示词已格式化，文档长度: {len(content)}")
            messages = [
                SystemMessage(content="你是一位资深的测试专家。"),
                HumanMessage(content=formatted_prompt)
            ]
            
            response = self.llm.invoke(messages)
            json_match = re.search(r'```json\s*(.*?)\s*```', response.content, re.DOTALL)
            
            if json_match:
                return json.loads(json_match.group(1))
            else:
                logger.warning("可测性分析未返回JSON格式，使用默认结果")
                return self._get_default_analysis_result('testability_analysis')
                
        except Exception as e:
            logger.error(f"可测性分析失败: {e}")
            import traceback
            logger.error(f"详细错误: {traceback.format_exc()}")
            return self._get_default_analysis_result('testability_analysis')
    
    def analyze_feasibility(self, content: str) -> dict:
        """可行性专项分析 - 分析完整文档的可行性"""
        feasibility_prompt = self._get_user_prompt('feasibility_analysis')
        if not feasibility_prompt:
            logger.warning("用户未配置可行性分析提示词，返回默认结果")
            return self._get_default_analysis_result('feasibility_analysis')
        
        try:
            prompt_template = Template(feasibility_prompt)
            formatted_prompt = prompt_template.safe_substitute(document=content)
            messages = [
                SystemMessage(content="你是一位资深的技术架构师。"),
                HumanMessage(content=formatted_prompt)
            ]
            
            response = self.llm.invoke(messages)
            json_match = re.search(r'```json\s*(.*?)\s*```', response.content, re.DOTALL)
            
            if json_match:
                return json.loads(json_match.group(1))
            else:
                logger.warning("可行性分析未返回JSON格式，使用默认结果")
                return self._get_default_analysis_result('feasibility_analysis')
                
        except Exception as e:
            logger.error(f"可行性分析失败: {e}")
            import traceback
            logger.error(f"详细错误: {traceback.format_exc()}")
            return self._get_default_analysis_result('feasibility_analysis')
    
    def analyze_clarity(self, content: str) -> dict:
        """清晰度专项分析 - 分析完整文档的清晰度"""
        clarity_prompt = self._get_user_prompt('clarity_analysis')
        if not clarity_prompt:
            logger.warning("用户未配置清晰度分析提示词，返回默认结果")
            return self._get_default_analysis_result('clarity_analysis')
        
        try:
            prompt_template = Template(clarity_prompt)
            formatted_prompt = prompt_template.safe_substitute(document=content)
            messages = [
                SystemMessage(content="你是一位资深的需求分析专家。"),
                HumanMessage(content=formatted_prompt)
            ]
            
            response = self.llm.invoke(messages)
            json_match = re.search(r'```json\s*(.*?)\s*```', response.content, re.DOTALL)
            
            if json_match:
                return json.loads(json_match.group(1))
            else:
                logger.warning("清晰度分析未返回JSON格式，使用默认结果")
                return self._get_default_analysis_result('clarity_analysis')
                
        except Exception as e:
            logger.error(f"清晰度分析失败: {e}")
            import traceback
            logger.error(f"详细错误: {traceback.format_exc()}")
            return self._get_default_analysis_result('clarity_analysis')
    
    def _get_default_analysis_result(self, analysis_type: str) -> dict:
        """获取默认的分析结果"""
        return {
            "analysis_type": analysis_type,
            "overall_score": 70,
            "summary": f"{analysis_type}分析完成，整体质量中等",
            "issues": []
        }

    def analyze_document_comprehensive(self, document: RequirementDocument, analysis_options: dict = None) -> dict:
        """
        全面分析需求文档 - 新架构：并发执行5个专项分析
        现在5个分析可以并发执行，提高效率
        
        Args:
            document: 要分析的文档
            analysis_options: 分析选项，可包含max_workers控制并发数
        """
        analysis_options = analysis_options or {}
        max_workers = analysis_options.get('max_workers', 3)  # 从选项中获取，默认3
        
        try:
            logger.info(f"开始全面分析文档: {document.title}, 内容长度: {len(document.content)}, 并发数: {max_workers}")
            
            # 使用线程池并发执行5个专项分析（每个都处理完整文档，充分利用200k上下文）
            from concurrent.futures import ThreadPoolExecutor, as_completed
            
            logger.info("开始并发执行5个专项分析...")
            
            # 定义5个分析任务
            analysis_tasks = {
                'completeness': ('完整性', self.analyze_completeness),
                'consistency': ('一致性', self.analyze_consistency),
                'testability': ('可测性', self.analyze_testability),
                'feasibility': ('可行性', self.analyze_feasibility),
                'clarity': ('清晰度', self.analyze_clarity),
            }
            
            # 并发执行所有分析
            results = {}
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                # 提交所有任务
                future_to_analysis = {
                    executor.submit(task_func, document.content): (name, display_name)
                    for name, (display_name, task_func) in analysis_tasks.items()
                }
                
                # 收集结果
                for future in as_completed(future_to_analysis):
                    analysis_name, display_name = future_to_analysis[future]
                    try:
                        result = future.result()
                        results[analysis_name] = result
                        logger.info(f"{display_name}分析完成，评分: {result.get('overall_score', 0)}")
                    except Exception as e:
                        logger.error(f"{display_name}分析失败: {e}")
                        # 使用默认结果
                        results[analysis_name] = self._get_default_analysis_result(f'{analysis_name}_analysis')
            
            logger.info("所有专项分析并发执行完成")
            
            # 生成综合报告（新版本）
            logger.info("生成综合分析报告...")
            comprehensive_report = self._generate_comprehensive_report_v2({
                'completeness': results.get('completeness', {}),
                'consistency': results.get('consistency', {}),
                'testability': results.get('testability', {}),
                'feasibility': results.get('feasibility', {}),
                'clarity': results.get('clarity', {}),
                'document': document
            })
            
            logger.info(f"文档分析完成，总体评分: {comprehensive_report.get('overall_score', 0)}")
            return comprehensive_report

        except Exception as e:
            logger.error(f"需求文档分析失败: {e}")
            import traceback
            logger.error(f"详细错误: {traceback.format_exc()}")
            raise
    
    def _generate_comprehensive_report_v2(self, analyses: dict) -> dict:
        """生成综合评审报告 - 新架构版本"""
        try:
            # 提取各专项分析结果
            completeness = analyses.get('completeness', {})
            consistency = analyses.get('consistency', {})
            testability = analyses.get('testability', {})
            feasibility = analyses.get('feasibility', {})
            clarity = analyses.get('clarity', {})
            document = analyses.get('document')
            
            # 计算总体评分（5个维度平均）
            scores = []
            for analysis in [completeness, consistency, testability, feasibility, clarity]:
                score = analysis.get('overall_score', 70)
                scores.append(score)
            
            overall_score = int(sum(scores) / len(scores)) if scores else 70
            
            # 收集所有问题
            all_issues = []
            for analysis_name, analysis_data in [
                ('completeness', completeness),
                ('consistency', consistency),
                ('testability', testability),
                ('feasibility', feasibility),
                ('clarity', clarity)
            ]:
                issues = analysis_data.get('issues', [])
                for issue in issues:
                    issue['source'] = analysis_name
                    issue['analysis_type'] = analysis_name
                    # 如果LLM返回的是severity而非priority,统一映射为priority
                    if 'severity' in issue and 'priority' not in issue:
                        issue['priority'] = issue['severity']
                all_issues.extend(issues)
            
            # 按优先级分类问题
            high_priority = [i for i in all_issues if i.get('priority') == 'high']
            medium_priority = [i for i in all_issues if i.get('priority') == 'medium']
            low_priority = [i for i in all_issues if i.get('priority') == 'low']
            
            logger.info(f"问题统计: 总数={len(all_issues)}, 高优先级={len(high_priority)}, 中优先级={len(medium_priority)}, 低优先级={len(low_priority)}")
            if all_issues and len(high_priority) == len(medium_priority) == len(low_priority) == 0:
                # 检查priority字段值
                priority_values = set(i.get('priority') for i in all_issues if 'priority' in i)
                logger.warning(f"所有问题都未分类到high/medium/low! 发现的priority值: {priority_values}")
            
            # 确定总体评价
            if overall_score >= 90:
                overall_rating = 'excellent'
            elif overall_score >= 80:
                overall_rating = 'good'
            elif overall_score >= 70:
                overall_rating = 'average'
            elif overall_score >= 60:
                overall_rating = 'needs_improvement'
            else:
                overall_rating = 'poor'
            
            # 收集改进建议
            recommendations = []
            for analysis in [completeness, consistency, testability, feasibility, clarity]:
                recs = analysis.get('recommendations', [])
                if isinstance(recs, list):
                    recommendations.extend(recs)
                elif isinstance(recs, str):
                    recommendations.append(recs)
            
            # 去重
            recommendations = list(set(recommendations))[:10]
            
            # 生成总结
            summary = self._generate_summary(overall_score, len(all_issues), len(high_priority))
            
            # 构建综合报告
            comprehensive_report = {
                'overall_score': overall_score,
                'overall_rating': overall_rating,
                'total_issues': len(all_issues),
                'high_priority_issues': len(high_priority),
                'medium_priority_issues': len(medium_priority),
                'low_priority_issues': len(low_priority),
                
                # 详细评分（5个专项维度）
                'scores': {
                    'completeness': completeness.get('overall_score', 70),
                    'consistency': consistency.get('overall_score', 70),
                    'testability': testability.get('overall_score', 70),
                    'feasibility': feasibility.get('overall_score', 70),
                    'clarity': clarity.get('overall_score', 70),
                },
                
                # 问题详情
                'issues': all_issues,
                'high_priority_issues_detail': high_priority,
                'medium_priority_issues_detail': medium_priority,
                'low_priority_issues_detail': low_priority,
                
                # 专项分析详情
                'specialized_analyses': {
                    'completeness_analysis': completeness,
                    'consistency_analysis': consistency,
                    'testability_analysis': testability,
                    'feasibility_analysis': feasibility,
                    'clarity_analysis': clarity
                },
                
                # 改进建议
                'recommendations': recommendations,
                
                # 总结
                'summary': summary,
                
                # 元数据
                'analysis_architecture': 'specialized_200k_context',
                'analysis_timestamp': str(document.updated_at) if document else '',
                'document_length': len(document.content) if document and document.content else 0
            }
            
            return comprehensive_report
            
        except Exception as e:
            logger.error(f"生成综合报告失败: {e}")
            import traceback
            logger.error(f"详细错误: {traceback.format_exc()}")
            raise

    def _analyze_global_structure(self, document: RequirementDocument) -> dict:
        """分析文档的全局结构和上下文"""

        # 从数据库获取用户的提示词
        global_prompt = self._get_user_prompt('global_analysis')
        if not global_prompt:
            raise ValueError("用户未配置全局分析提示词，请先在提示词管理中配置")

        try:
            prompt_template = Template(global_prompt)
            formatted_prompt = prompt_template.safe_substitute(
                title=document.title,
                description=document.description or "无描述",
                content=document.content[:6000]
            )
            messages = [
                SystemMessage(content="你是一位专业的需求分析师，擅长需求文档评审。"),
                HumanMessage(content=formatted_prompt)
            ]

            response = self.llm.invoke(messages)

            # 提取JSON内容
            json_match = re.search(r'```json\s*(.*?)\s*```', response.content, re.DOTALL)
            if json_match:
                json_str = json_match.group(1)
                global_analysis = json.loads(json_str)
            else:
                # 如果没有找到JSON，返回默认结构
                global_analysis = self._get_default_global_analysis()

            logger.info(f"全局分析完成，总体评分: {global_analysis.get('overall_score', 0)}")
            return global_analysis

        except json.JSONDecodeError as e:
            logger.error(f"解析全局分析JSON失败: {e}")
            return self._get_default_global_analysis()
        except Exception as e:
            logger.error(f"全局结构分析失败: {e}")
            return self._get_default_global_analysis()

    def _analyze_modules_detailed(self, document: RequirementDocument, global_context: dict) -> List[dict]:
        """详细分析各个模块"""
        modules = document.modules.order_by('order')
        module_analyses = []

        for module in modules:
            try:
                analysis = self._analyze_single_module(module, global_context)
                module_analyses.append(analysis)
            except Exception as e:
                logger.error(f"模块 {module.title} 分析失败: {e}")
                # 添加默认分析结果
                module_analyses.append(self._get_default_module_analysis(module))

        return module_analyses

    def _analyze_single_module(self, module: RequirementModule, global_context: dict) -> dict:
        """分析单个模块"""

        # 从数据库获取用户的提示词
        module_prompt = self._get_user_prompt('module_analysis')
        if not module_prompt:
            raise ValueError("用户未配置模块分析提示词，请先在提示词管理中配置")

        try:
            prompt_template = Template(module_prompt)
            formatted_prompt = prompt_template.safe_substitute(
                module_id=str(module.id),
                module_title=module.title,
                module_content=module.content[:3000],
                business_flows=", ".join(global_context.get('business_flows', [])),
                data_entities=", ".join(global_context.get('data_entities', [])),
                global_rules=", ".join(global_context.get('global_rules', []))
            )
            messages = [
                SystemMessage(content="你是一位专业的需求分析师，正在进行需求评审。"),
                HumanMessage(content=formatted_prompt)
            ]

            response = self.llm.invoke(messages)

            # 提取JSON内容
            json_match = re.search(r'```json\s*(.*?)\s*```', response.content, re.DOTALL)
            if json_match:
                json_str = json_match.group(1)
                analysis = json.loads(json_str)
                analysis['module_id'] = str(module.id)  # 确保ID正确
                return analysis
            else:
                return self._get_default_module_analysis(module)

        except json.JSONDecodeError as e:
            logger.error(f"解析模块分析JSON失败: {e}")
            return self._get_default_module_analysis(module)
        except Exception as e:
            logger.error(f"模块分析失败: {e}")
            return self._get_default_module_analysis(module)

    def _analyze_cross_module_consistency(self, document: RequirementDocument,
                                        module_analyses: List[dict], global_context: dict) -> dict:
        """分析跨模块一致性"""

        # 从数据库获取用户的提示词
        consistency_prompt = self._get_user_prompt('consistency_analysis')
        if not consistency_prompt:
            raise ValueError("用户未配置一致性分析提示词，请先在提示词管理中配置")

        try:
            # 准备上下文数据
            context_str = json.dumps(global_context, ensure_ascii=False, indent=2)
            analyses_str = json.dumps(module_analyses, ensure_ascii=False, indent=2)

            prompt_template = Template(consistency_prompt)
            formatted_prompt = prompt_template.safe_substitute(
                global_context=context_str[:2000],
                module_analyses=analyses_str[:4000]
            )
            messages = [
                SystemMessage(content="你是一位专业的需求分析师，擅长跨模块一致性检查。"),
                HumanMessage(content=formatted_prompt)
            ]

            response = self.llm.invoke(messages)

            # 提取JSON内容
            json_match = re.search(r'```json\s*(.*?)\s*```', response.content, re.DOTALL)
            if json_match:
                json_str = json_match.group(1)
                consistency_analysis = json.loads(json_str)
            else:
                consistency_analysis = self._get_default_consistency_analysis()

            logger.info(f"一致性分析完成，一致性评分: {consistency_analysis.get('consistency_score', 0)}")
            return consistency_analysis

        except json.JSONDecodeError as e:
            logger.error(f"解析一致性分析JSON失败: {e}")
            return self._get_default_consistency_analysis()
        except Exception as e:
            logger.error(f"一致性分析失败: {e}")
            return self._get_default_consistency_analysis()

    def _generate_comprehensive_report(self, global_analysis: dict,
                                     module_analyses: List[dict],
                                     consistency_analysis: dict) -> dict:
        """生成综合评审报告"""

        # 计算总体评分
        global_score = global_analysis.get('overall_score', 0)
        module_scores = [m.get('overall_score', 0) for m in module_analyses]
        avg_module_score = sum(module_scores) / len(module_scores) if module_scores else 0
        consistency_score = consistency_analysis.get('consistency_score', 0)

        overall_score = int((global_score * 0.3 + avg_module_score * 0.5 + consistency_score * 0.2))

        # 统计问题数量
        all_issues = []

        # 收集模块问题
        for module_analysis in module_analyses:
            module_issues = module_analysis.get('issues', [])
            for issue in module_issues:
                issue['source'] = 'module'
                issue['module_name'] = module_analysis.get('module_name', '未知模块')
            all_issues.extend(module_issues)

        # 收集一致性问题
        consistency_issues = consistency_analysis.get('cross_module_issues', [])
        for issue in consistency_issues:
            issue['source'] = 'consistency'
        all_issues.extend(consistency_issues)

        # 按优先级分类问题
        high_priority = [i for i in all_issues if i.get('priority') == 'high']
        medium_priority = [i for i in all_issues if i.get('priority') == 'medium']
        low_priority = [i for i in all_issues if i.get('priority') == 'low']

        # 确定总体评价
        if overall_score >= 90:
            overall_rating = 'excellent'
        elif overall_score >= 80:
            overall_rating = 'good'
        elif overall_score >= 70:
            overall_rating = 'average'
        elif overall_score >= 60:
            overall_rating = 'needs_improvement'
        else:
            overall_rating = 'poor'

        # 生成改进建议
        recommendations = []
        recommendations.extend(global_analysis.get('missing_aspects', []))

        for module_analysis in module_analyses:
            recommendations.extend(module_analysis.get('recommendations', []))

        recommendations.extend(consistency_analysis.get('recommendations', []))

        # 去重
        recommendations = list(set(recommendations))

        comprehensive_report = {
            'overall_score': overall_score,
            'overall_rating': overall_rating,
            'total_issues': len(all_issues),
            'high_priority_issues': len(high_priority),
            'medium_priority_issues': len(medium_priority),
            'low_priority_issues': len(low_priority),

            # 详细评分
            'scores': {
                'global_structure': global_score,
                'module_average': int(avg_module_score),
                'consistency': consistency_score,
                'specification': int(sum(m.get('specification_score', 0) for m in module_analyses) / len(module_analyses)) if module_analyses else 0,
                'clarity': int(sum(m.get('clarity_score', 0) for m in module_analyses) / len(module_analyses)) if module_analyses else 0,
                'completeness': int(sum(m.get('completeness_score', 0) for m in module_analyses) / len(module_analyses)) if module_analyses else 0,
                'feasibility': int(sum(m.get('feasibility_score', 0) for m in module_analyses) / len(module_analyses)) if module_analyses else 0
            },

            # 问题详情
            'issues': all_issues,
            'high_priority_issues_detail': high_priority,
            'medium_priority_issues_detail': medium_priority,
            'low_priority_issues_detail': low_priority,

            # 分析详情
            'global_analysis': global_analysis,
            'module_analyses': module_analyses,
            'consistency_analysis': consistency_analysis,

            # 改进建议
            'recommendations': recommendations[:10],  # 限制数量

            # 总结
            'summary': self._generate_summary(overall_score, len(all_issues), len(high_priority)),

            # 元数据
            'analysis_timestamp': logger.info.__self__.__class__.__name__,  # 简单的时间戳
            'modules_count': len(module_analyses)
        }

        return comprehensive_report

    def _generate_summary(self, overall_score: int, total_issues: int, high_priority_issues: int) -> str:
        """生成评审总结"""
        if overall_score >= 90:
            quality = "优秀"
        elif overall_score >= 80:
            quality = "良好"
        elif overall_score >= 70:
            quality = "一般"
        elif overall_score >= 60:
            quality = "需要改进"
        else:
            quality = "较差"

        summary = f"需求文档整体质量{quality}（评分：{overall_score}/100）。"

        if high_priority_issues > 0:
            summary += f"发现{high_priority_issues}个高优先级问题需要立即解决。"

        if total_issues > 10:
            summary += f"共发现{total_issues}个问题，建议优先处理高优先级问题。"
        elif total_issues > 0:
            summary += f"发现{total_issues}个问题，整体可控。"
        else:
            summary += "未发现明显问题，质量较高。"

        return summary

    def _get_default_global_analysis(self) -> dict:
        """获取默认的全局分析结果"""
        return {
            "structure_score": 70,
            "completeness_score": 70,
            "consistency_score": 70,
            "clarity_score": 70,
            "overall_score": 70,
            "business_flows": [],
            "data_entities": [],
            "global_rules": [],
            "missing_aspects": ["需要更详细的分析"],
            "risk_points": [],
            "strengths": ["文档结构基本合理"],
            "weaknesses": ["需要更详细的需求描述"]
        }

    def _get_default_module_analysis(self, module: RequirementModule) -> dict:
        """获取默认的模块分析结果"""
        return {
            "module_id": str(module.id),
            "module_name": module.title,
            "specification_score": 70,
            "clarity_score": 70,
            "completeness_score": 70,
            "consistency_score": 70,
            "feasibility_score": 70,
            "overall_score": 70,
            "issues": [],
            "strengths": ["模块功能基本清晰"],
            "weaknesses": ["需要更详细的描述"],
            "recommendations": ["补充详细的功能说明"]
        }

    def _get_default_consistency_analysis(self) -> dict:
        """获取默认的一致性分析结果"""
        return {
            "consistency_score": 70,
            "interface_consistency": 70,
            "data_consistency": 70,
            "business_rule_consistency": 70,
            "process_completeness": 70,
            "cross_module_issues": [],
            "missing_connections": [],
            "redundant_functions": [],
            "recommendations": ["建议进行更详细的一致性检查"]
        }


class RequirementReviewService:
    """需求评审服务 - 统一的评审管理服务"""

    def __init__(self, user=None):
        self.user = user
        self.review_engine = RequirementReviewEngine(user=user)

    def start_direct_review(self, document: RequirementDocument,
                          analysis_options: dict = None) -> 'ReviewReport':
        """启动直接评审（不拆分模块）"""
        from .models import ReviewReport, ReviewIssue

        try:
            # 检查文档内容
            if not document.content:
                raise ValueError("文档内容为空，无法进行评审")

            # 创建评审报告
            review_report = ReviewReport.objects.create(
                document=document,
                status='in_progress',
                reviewer='AI需求评审助手',
                review_type='direct'  # 标记为直接评审
            )

            # 更新文档状态
            document.status = 'reviewing'
            document.save()

            # 对整个文档进行评审
            review_result = self.review_engine.analyze_document_directly(
                document.content,
                analysis_options or {}
            )

            # 更新评审报告
            review_report.overall_rating = review_result.get('overall_rating', 'average')
            review_report.completion_score = review_result.get('completion_score', 0)
            review_report.clarity_score = review_result.get('clarity_score', 0)
            review_report.consistency_score = review_result.get('consistency_score', 0)
            review_report.completeness_score = review_result.get('completeness_score', 0)
            review_report.summary = review_result.get('summary', '')
            review_report.recommendations = review_result.get('recommendations', '')
            review_report.status = 'completed'
            review_report.save()

            # 创建问题记录
            issues = review_result.get('issues', [])
            for issue_data in issues:
                # 映射字段名
                issue_params = {
                    'report': review_report,
                    'title': issue_data.get('title', '未知问题'),
                    'description': issue_data.get('description', ''),
                    'priority': issue_data.get('priority', 'medium'),
                    'issue_type': issue_data.get('category', 'specification'),  # category -> issue_type
                    'suggestion': issue_data.get('suggestion', ''),
                    'location': issue_data.get('location', '')
                }
                ReviewIssue.objects.create(**issue_params)

            # 更新统计信息
            review_report.total_issues = len(issues)
            review_report.high_priority_issues = len([i for i in issues if i.get('priority') == 'high'])
            review_report.save()

            # 更新文档状态
            document.status = 'review_completed'
            document.save()

            logger.info(f"文档 {document.id} 直接评审完成")
            return review_report

        except Exception as e:
            logger.error(f"直接评审失败: {e}")
            if 'review_report' in locals():
                review_report.status = 'failed'
                review_report.save()
            raise

    def start_comprehensive_review(self, document: RequirementDocument,
                                 analysis_options: dict = None) -> 'ReviewReport':
        """启动全面的需求评审（基于模块）"""
        from .models import ReviewReport, ReviewIssue, ModuleReviewResult

        try:
            # 检查文档状态（允许 ready_for_review 或 reviewing 状态）
            if document.status not in ['ready_for_review', 'reviewing']:
                raise ValueError(f"文档状态 {document.status} 不允许开始评审")

            # 创建评审报告
            review_report = ReviewReport.objects.create(
                document=document,
                status='in_progress',
                reviewer='AI需求评审助手',
                review_type='comprehensive'  # 标记为全面评审
            )

            # 确保文档状态为 reviewing
            if document.status != 'reviewing':
                document.status = 'reviewing'
                document.save()

            logger.info(f"开始评审文档: {document.title}")

            # 执行AI分析
            analysis_result = self.review_engine.analyze_document_comprehensive(document, analysis_options)

            # 更新评审报告
            self._update_review_report(review_report, analysis_result)

            # 创建问题记录
            self._create_review_issues(review_report, analysis_result)

            # 创建模块评审结果
            self._create_module_results(review_report, analysis_result)

            # 完成评审
            review_report.status = 'completed'
            review_report.save()

            # 更新文档状态
            document.status = 'review_completed'
            document.save()

            logger.info(f"评审完成: {document.title}, 总体评分: {review_report.completion_score}")

            return review_report

        except Exception as e:
            logger.error(f"评审失败: {e}")

            # 更新失败状态
            if 'review_report' in locals():
                review_report.status = 'failed'
                review_report.save()

            document.status = 'failed'
            document.save()

            raise

    def _update_review_report(self, review_report: 'ReviewReport', analysis_result: dict):
        """更新评审报告基本信息和专项分析详情"""
        review_report.overall_rating = analysis_result.get('overall_rating', 'average')
        review_report.completion_score = analysis_result.get('overall_score', 0)
        review_report.total_issues = analysis_result.get('total_issues', 0)
        review_report.high_priority_issues = analysis_result.get('high_priority_issues', 0)
        review_report.medium_priority_issues = analysis_result.get('medium_priority_issues', 0)
        review_report.low_priority_issues = analysis_result.get('low_priority_issues', 0)
        review_report.summary = analysis_result.get('summary', '')
        review_report.recommendations = '\n'.join(analysis_result.get('recommendations', []))
        
        # 保存专项分析详情（包含issues, strengths, recommendations等完整数据）
        specialized_analyses = analysis_result.get('specialized_analyses', {})
        review_report.specialized_analyses = specialized_analyses
        
        # 同时保存各专项分析的分数到独立字段
        review_report.completeness_score = specialized_analyses.get('completeness_analysis', {}).get('overall_score', 0)
        review_report.consistency_score = specialized_analyses.get('consistency_analysis', {}).get('overall_score', 0)
        review_report.clarity_score = specialized_analyses.get('clarity_analysis', {}).get('overall_score', 0)
        review_report.testability_score = specialized_analyses.get('testability_analysis', {}).get('overall_score', 0)
        review_report.feasibility_score = specialized_analyses.get('feasibility_analysis', {}).get('overall_score', 0)
        
        review_report.save()

    def _create_review_issues(self, review_report: 'ReviewReport', analysis_result: dict):
        """创建评审问题记录"""
        from .models import ReviewIssue

        issues = analysis_result.get('issues', [])

        for issue_data in issues:
            try:
                # 查找相关模块
                module = None
                if issue_data.get('module_name'):
                    module = review_report.document.modules.filter(
                        title__icontains=issue_data['module_name']
                    ).first()

                # 映射问题类型
                issue_type = self._map_issue_type(issue_data.get('type', 'clarity'))

                ReviewIssue.objects.create(
                    report=review_report,
                    module=module,
                    issue_type=issue_type,
                    priority=issue_data.get('priority', 'medium'),
                    title=issue_data.get('title', '未知问题'),
                    description=issue_data.get('description', ''),
                    suggestion=issue_data.get('suggestion', ''),
                    location=issue_data.get('location', ''),
                    section=issue_data.get('module_name', '')
                )

            except Exception as e:
                logger.error(f"创建问题记录失败: {e}")

    def _create_module_results(self, review_report: 'ReviewReport', analysis_result: dict):
        """创建模块评审结果"""
        from .models import ModuleReviewResult

        module_analyses = analysis_result.get('module_analyses', [])

        for module_analysis in module_analyses:
            try:
                # 查找模块
                module_id = module_analysis.get('module_id')
                if not module_id:
                    continue

                module = review_report.document.modules.filter(id=module_id).first()
                if not module:
                    continue

                # 计算严重程度评分（分数越高问题越严重）
                overall_score = module_analysis.get('overall_score', 70)
                severity_score = max(0, 100 - overall_score)

                # 映射评级
                module_rating = self._map_module_rating(overall_score)

                ModuleReviewResult.objects.create(
                    report=review_report,
                    module=module,
                    module_rating=module_rating,
                    issues_count=len(module_analysis.get('issues', [])),
                    severity_score=severity_score,
                    analysis_content=json.dumps(module_analysis, ensure_ascii=False, indent=2),
                    strengths='\n'.join(module_analysis.get('strengths', [])),
                    weaknesses='\n'.join(module_analysis.get('weaknesses', [])),
                    recommendations='\n'.join(module_analysis.get('recommendations', []))
                )

            except Exception as e:
                logger.error(f"创建模块结果失败: {e}")

    def _map_issue_type(self, ai_type: str) -> str:
        """映射AI分析的问题类型到数据库字段"""
        type_mapping = {
            'specification': 'specification',
            'clarity': 'clarity',
            'completeness': 'completeness',
            'consistency': 'consistency',
            'feasibility': 'feasibility',
            'data_inconsistency': 'consistency',
            'interface_inconsistency': 'consistency',
            'business_rule_inconsistency': 'consistency'
        }
        return type_mapping.get(ai_type, 'clarity')

    def _map_module_rating(self, score: int) -> str:
        """映射评分到评级"""
        if score >= 90:
            return 'excellent'
        elif score >= 80:
            return 'good'
        elif score >= 70:
            return 'average'
        elif score >= 60:
            return 'needs_improvement'
        else:
            return 'poor'

    def get_review_progress(self, document: RequirementDocument) -> dict:
        """获取评审进度（模拟实现）"""
        # 这里可以实现真实的进度跟踪
        # 目前返回模拟数据

        latest_review = document.review_reports.order_by('-review_date').first()
        if not latest_review:
            return {
                'status': 'not_started',
                'progress': 0,
                'message': '尚未开始评审'
            }

        if latest_review.status == 'completed':
            return {
                'status': 'completed',
                'progress': 100,
                'message': '评审已完成',
                'report_id': str(latest_review.id)
            }
        elif latest_review.status == 'in_progress':
            return {
                'status': 'in_progress',
                'progress': 75,  # 模拟进度
                'message': '正在进行评审分析...',
                'current_step': '分析模块一致性',
                'estimated_remaining': '2分钟'
            }
        else:
            return {
                'status': 'failed',
                'progress': 0,
                'message': '评审失败，请重试'
            }
